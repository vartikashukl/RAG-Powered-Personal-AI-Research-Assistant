# utils/loaders.py
import os
from typing import List
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders import WikipediaLoader

def load_pdf(file_path: str) -> List[Document]:
    """Load PDF file and return documents."""
    try:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        # Add source metadata
        for doc in documents:
            doc.metadata["source"] = file_path
        return documents
    except Exception as e:
        print(f"❌ Error loading PDF {file_path}: {e}")
        return []

def load_txt(file_path: str) -> List[Document]:
    """Load text file and return documents."""
    try:
        loader = TextLoader(file_path, encoding='utf-8')
        documents = loader.load()
        # Add source metadata
        for doc in documents:
            doc.metadata["source"] = file_path
        return documents
    except Exception as e:
        print(f"❌ Error loading TXT {file_path}: {e}")
        return []

def load_docx(file_path: str) -> List[Document]:
    """Load DOCX file and return documents."""
    try:
        # Simple DOCX loader using python-docx
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        document = Document(
            page_content=text,
            metadata={"source": file_path}
        )
        return [document]
    except Exception as e:
        print(f"❌ Error loading DOCX {file_path}: {e}")
        return []

def load_wikipedia(topic: str, max_docs: int = 3) -> List[Document]:
    """Load Wikipedia articles on a topic."""
    try:
        loader = WikipediaLoader(query=topic, load_max_docs=max_docs)
        documents = loader.load()
        # Add source metadata
        for doc in documents:
            doc.metadata["source"] = f"Wikipedia: {topic}"
        return documents
    except Exception as e:
        print(f"❌ Error loading Wikipedia topic '{topic}': {e}")
        return []